from pathlib import Path
import re
import shutil

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)

BASE = Path(r"c:\Users\fanele\AI-Driven Controller Placement Optimization in Multi-Site Software Defined Networks\ai-sdn-controller-placement\docs")
DOCX_PATH = BASE / "Research Proposal Template - PG Studies - Hons.docx"
PDF_PATH = BASE / "Research Proposal  - PG Studies - Hons.pdf"
TEMPLATE_PDF_PATH = BASE / "Research Proposal Template - PG Studies - Hons.pdf"

TITLE = "Evidence-Based Multi-Objective Evaluation of AI and Heuristic Controller Placement for Multi-Site Software Defined Networks"

REFERENCE_ITEMS = [
    (
        "1",
        "B. Heller, R. Sherwood, and N. McKeown, \"The controller placement problem,\" Proc. HotSDN, pp. 7-12, 2012.",
        "https://doi.org/10.1145/2342441.2342444",
    ),
    (
        "2",
        "T. F. Gonzalez, \"Clustering to minimize the maximum intercluster distance,\" Theoretical Computer Science, vol. 38, pp. 293-306, 1985.",
        "https://doi.org/10.1016/0304-3975(85)90224-5",
    ),
    (
        "3",
        "J. H. Holland, Adaptation in Natural and Artificial Systems. University of Michigan Press, 1975.",
        "https://books.google.com/books?id=JE5RAAAAMAAJ",
    ),
    (
        "4",
        "J. MacQueen, \"Some methods for classification and analysis of multivariate observations,\" Proc. Fifth Berkeley Symp., vol. 1, pp. 281-297, 1967.",
        "https://projecteuclid.org/euclid.bsmsp/1200512992",
    ),
    (
        "5",
        "R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed. MIT Press, 2018.",
        "https://www.andrew.cmu.edu/course/10-703/textbook/BartoSutton.pdf",
    ),
    (
        "6",
        "M. Benoudifa et al., \"MuZero-based autonomous controller placement,\" IJCNIS, vol. 15, no. 1, pp. 1-18, 2023.",
        "https://ijcnis.org/index.php/ijcnis/article/view/6787",
    ),
    (
        "7",
        "A. Farahi et al., \"A comprehensive review on SDN controller placement problem and new opportunities in AI and optimization methods,\" Computers, vol. 11, no. 8, 2022.",
        "https://doi.org/10.3390/computers11080111",
    ),
    (
        "8",
        "M. Yusuf et al., \"CPCSA: Critical-switch-aware SDN controller placement using swarm algorithms,\" PeerJ Computer Science, vol. 9, e1698, 2023.",
        "https://doi.org/10.7717/peerj-cs.1698",
    ),
    (
        "9",
        "D. Kreutz et al., \"Software-defined networking: A comprehensive survey,\" Proceedings of the IEEE, vol. 103, no. 1, pp. 14-76, 2015.",
        "https://doi.org/10.1109/JPROC.2014.2371999",
    ),
]

REFERENCE_IDS = {item[0] for item in REFERENCE_ITEMS}
CITATION_PATTERN = re.compile(r"\[(\d+)\]")

sections = [
    ("1. Introduction", [
        "Software Defined Networking (SDN) controller placement is a critical infrastructure optimisation problem because controller location and controller count influence latency, reliability, and operational efficiency. This study focuses on whether AI-driven controller placement methods provide decision-relevant advantages over heuristic methods in multi-site SDN environments. The topic is important because network planners need defensible, evidence-based guidance when selecting methods for performance-sensitive and resource-constrained deployments.",
        "This proposal is structured as follows. Section 2 presents the background to SDN controller placement and identifies under-investigated areas in current research. Section 3 defines the problem statement and research questions. Section 4 states the research aim and objectives, followed by Section 5 on contribution to knowledge and practice. Sections 6 to 10 cover the literature review, conceptual framework, methodology, ethical and safety issues, and intellectual property. Sections 11 to 15 present the resources required, feasibility, dissemination, dissertation structure, and referencing approach.",
    ]),
    ("2. Background of the Study", [
        "Controller placement in SDN is commonly formulated as an optimisation problem that assigns switches to controllers under latency and resilience constraints. The topic is important because controller decisions directly affect control-plane responsiveness, continuity under failures, and overall network manageability, especially in multi-site environments where topology and demand may vary over time.",
        "Current research shows two broad method families. Classical heuristics, including greedy and clustering-based approaches, are computationally efficient and interpretable. AI-oriented methods, including genetic algorithms and reinforcement learning, are reported to improve placement quality in selected scenarios. However, important gaps remain: many studies prioritise single outcomes such as latency while under-reporting runtime and convergence costs; cross-topology and cross-scale robustness is often insufficiently tested; and reproducibility protocols are not always explicit.",
        "This study addresses these gaps through a reproducible, multi-objective evaluation pipeline that compares AI and non-AI methods under controlled factorial scenarios. It advances knowledge by producing transparent evidence on trade-offs among latency, reliability, and computational cost across topology families and scales. It advances practice by generating scenario-conditioned guidance on when AI methods are justified and when heuristic methods are more suitable.",
    ]),
    ("3. Problem Statement and Research Questions", []),
    ("3.1 Problem Statement", [
        "There is insufficient reproducible, multi-objective evidence to determine when AI-based SDN controller placement methods outperform classical heuristics in multi-site environments, creating uncertainty in operational method selection.",
    ]),
    ("3.2 Research Questions", [
        "1. In multi-site SDN topologies, do AI-based controller placement methods significantly outperform classical heuristics on control-plane latency?",
        "2. How do AI and heuristic methods compare on reliability and computational cost across topology families and network scales?",
        "3. Under which scenario conditions is the additional complexity of AI methods justified over heuristic baselines?",
    ]),
    ("4. Research Aim and Objectives", []),
    ("4.1 Research Aim", [
        "To evaluate and compare the trade-offs between AI-driven and heuristic-based SDN controller placement methods in multi-site software-defined networks using a reproducible, multi-objective factorial experimental design.",
    ]),
    ("4.2 Research Objectives", [
        "1. To compare AI-driven and heuristic-based controller placement methods on control-plane latency, reliability, and computational cost across multi-site SDN scenarios.",
        "2. To assess the robustness and generalisability of both method families across topology types and network scales using a controlled factorial experiment.",
        "3. To formulate scenario-conditioned decision guidelines that indicate when AI-based methods provide sufficient benefit over heuristic baselines to justify additional computational complexity.",
    ]),
    ("5. Contribution of the Study", []),
    ("5.1 Contribution to Knowledge", [
        "This study contributes to knowledge by developing a reproducible, multi-objective evaluation framework for SDN controller placement that integrates latency, reliability, and computational cost within one coherent analytical pipeline. It extends current understanding by moving beyond single-metric performance reporting and by explicitly modelling trade-offs across topology families and scale conditions. The study also strengthens methodological rigour in this domain through transparent baseline comparisons, controlled factorial design, and explicit evidence for scenario-dependent method performance.",
    ]),
    ("5.2 Contribution to Practice and Policy", [
        "This study contributes to practice by producing decision-ready guidance for infrastructure engineers and network planners on selecting controller placement methods under operational constraints. Rather than assuming AI superiority, it provides scenario-specific recommendations on when heuristic methods are sufficient and when AI methods deliver meaningful gains relative to cost. For institutional and organisational policy, the study supports adoption of reproducibility standards, benchmark transparency, and multi-criteria evaluation protocols in AI-enabled network optimisation projects.",
    ]),
    ("6. Literature Review", [
        "Software-defined networking controller placement has been studied extensively as a network optimisation problem in which controller location and controller count shape latency, reliability, and operational efficiency. The literature shows movement from classical heuristic methods toward metaheuristic and learning-based approaches, but the empirical evidence remains uneven in how it measures performance, reports scalability, and supports decision-making under real deployment constraints. This review is organised thematically according to the three research objectives: comparative performance, robustness across topology and scale, and decision frameworks for controller selection.",
    ]),
    ("6.1 Comparative Performance of AI and Heuristic Controller Placement Methods", [
        "Recent empirical studies consistently show that controller placement quality is not determined by a single metric. Classical heuristics such as greedy selection, clustering-based approaches, and k-center variants are often reported to deliver strong baseline performance because they are computationally efficient and easy to implement. In several experiments, these methods achieve competitive latency outcomes, particularly in small to medium networks, while keeping runtime overhead low. This makes them practical benchmarks for any proposal that claims improvement through more advanced methods [1][2][4][7].",
        "In contrast, AI-driven methods, including genetic algorithms and reinforcement learning, are frequently evaluated as potential improvements over heuristics. Empirical results often indicate that these methods can reduce average latency or improve controller coverage under selected conditions. However, the gains are rarely universal. Some studies report marginal improvements that depend on topology shape, controller count, or initialisation settings, while others show that the computational cost of training or iterative search offsets the performance gain. This is important because it suggests that superiority claims must be interpreted in relation to the cost of producing the result, not only the result itself [5][6][8].",
        "A recurring weakness in the literature is that latency is often treated as the dominant outcome, while reliability and runtime are under-emphasised. In studies where reliability is included, it is frequently operationalised in narrow ways, such as link failure tolerance or controller reachability, without a full multi-objective comparison against computational cost. This limitation matters for the present study because the research problem is not simply whether AI methods can outperform heuristics on one metric, but whether they provide a meaningful overall advantage once performance, reliability, and cost are considered together [7][9].",
    ]),
    ("6.2 Robustness Across Topology Families and Network Scales", [
        "A second major theme in the literature concerns robustness. Empirical SDN controller placement studies frequently test methods on a limited set of synthetic topologies, commonly small graphs or a single topology family, which restricts the generalisability of findings. Some studies show that a method performing well on one topology family may degrade sharply when evaluated on a different structural pattern, such as denser graphs, irregular multi-site layouts, or larger-scale deployments. This is especially relevant for multi-site environments, where controller placement must remain effective under more heterogeneous and distributed conditions [1][7][9].",
        "Recent work has increasingly recognised scale as a critical factor. Studies that expand experiments across network sizes often find that algorithm behaviour changes as the number of switches, links, and controllers increases. Heuristic methods tend to remain stable in runtime, while AI-based methods may exhibit better solution quality in some larger settings but at the cost of slower convergence or greater sensitivity to parameter choices. In other cases, the advantage of AI methods disappears as scale increases because the optimisation search becomes more expensive or less consistent across runs. These findings support the view that scalability is not a secondary detail but a central criterion for evaluating practical usefulness [6][7][8].",
        "The literature also shows that robustness is usually underreported in a way that weakens comparability. Many papers present results for a single topology family, a narrow range of scale levels, or a single experimental seed. Few studies report variation across multiple topology types together with performance dispersion or statistical stability. As a result, it is difficult to determine whether reported gains are reproducible or merely artefacts of a specific test setup. This directly matches the gap identified in the background section: cross-topology and cross-scale behaviour is still insufficiently documented [7][9].",
    ]),
    ("6.3 Decision Frameworks for Multi-Objective Controller Placement", [
        "A third body of literature addresses multi-objective decision-making in controller placement. The theoretical basis here is that controller placement must balance competing objectives rather than optimise one dimension in isolation. In principle, this aligns well with Pareto-based reasoning and trade-off analysis, where a solution is judged by how it performs across several criteria simultaneously. In practice, however, many empirical studies still stop at reporting performance indicators without translating those findings into explicit decision guidance [7][9].",
        "Several recent studies use multi-objective evolutionary methods or reinforcement learning variants to generate sets of candidate solutions rather than a single optimal solution. These studies show that trade-off frontiers can help reveal the relationship between latency reduction and other costs, such as runtime or resilience loss. Where Pareto analysis is used well, it helps demonstrate that the best method depends on the decision context rather than on absolute performance alone. This is useful for infrastructure planning because decision-makers need to know when a more complex method is worth the extra computational burden [5][6][8].",
        "Despite this progress, the literature still lacks a consistent decision framework for selecting among controller placement methods. Studies often present comparative tables or plots, but do not define practical thresholds for when AI methods justify their added complexity. Very few studies convert experimental outcomes into scenario-conditioned guidance such as: use the heuristic when the topology is small or static, and use the AI method only when expected performance gains exceed a defined cost threshold. That omission is significant because it leaves practitioners without a usable basis for method selection [7][9].",
    ]),
    ("6.4 Synthesis of the Literature and Research Gap", [
        "Taken together, the empirical literature shows that AI-driven controller placement can offer performance gains, but those gains are neither uniform nor automatically superior to heuristic baselines. Classical methods remain competitive in efficiency and predictability, while AI methods sometimes improve solution quality at the expense of runtime, stability, or implementation complexity. The strongest recurring limitation in current research is the lack of studies that compare these methods using a reproducible, multi-objective framework across multiple topology families and scales [1][6][7][8].",
        "This creates three under-investigated areas directly relevant to the present study. First, the literature does not yet provide enough evidence on comparative performance when latency, reliability, and computational cost are evaluated together. Second, cross-topology and cross-scale robustness is still insufficiently established. Third, existing studies rarely convert experimental results into practical decision rules for selecting controller placement methods in real multi-site SDN environments.",
        "Accordingly, this study is necessary because it addresses these gaps through a controlled empirical evaluation that compares AI-driven and heuristic methods using a consistent experimental design and multi-objective analysis. In line with the background section, the study will generate new knowledge about trade-offs, improve understanding of method behaviour under different conditions, and support better practice through scenario-based guidance for controller placement decisions [1][7][9].",
    ]),
    ("7. Theoretical or Conceptual Framework", [
        "This study is situated within three complementary concepts: optimisation theory, multi-objective decision analysis, and distributed systems reliability. Together, these concepts provide the explanatory basis for evaluating whether AI-driven controller placement methods justify their added complexity when compared with heuristic alternatives in multi-site SDN environments.",
        "At the core of the study is optimisation theory, which treats controller placement as a constrained decision problem in which feasible solutions must satisfy performance and resilience requirements. In this context, the goal is not only to identify a high-performing placement, but to determine how well each method navigates competing constraints such as latency, reliability, and computational cost. Heuristic methods are conceptually appropriate as efficient search strategies that produce acceptable solutions with low overhead, while AI-driven methods represent more complex search and learning strategies that may improve solution quality but require greater computational effort.",
        "The study is further anchored in multi-objective decision analysis, particularly the notion that no single controller placement method can be evaluated meaningfully on one metric alone. Instead, methods should be assessed as trade-off structures across multiple objectives. This framework supports comparison through dominance, balance, and Pareto efficiency, allowing the study to determine whether AI-based approaches offer materially better overall outcomes than heuristics when latency improvement is weighed against runtime burden and reliability behaviour. This is the key conceptual basis for judging whether AI complexity is justified over heuristic simplicity.",
        "Distributed systems reliability provides the third conceptual lens. In multi-site SDN deployments, controller placement is not merely an optimisation exercise; it is also a resilience problem because controller location affects fault tolerance, coverage under failure, and continuity of network control. This concept explains why the study includes reliability as a central performance dimension rather than treating it as a secondary outcome. It also supports the expectation that a method that performs well under ideal conditions may not remain effective when network topology, scale, or failure conditions change.",
        "Taken together, these concepts explain the anticipated findings of the study. The framework predicts that the comparative value of AI and heuristic placement methods will depend on the interaction between topology characteristics, scale, and the relative weighting of performance objectives. AI methods are expected to be beneficial only where their performance gains are sufficiently large to justify their computational and implementation cost. Heuristic methods are expected to remain competitive where efficiency, interpretability, and stability are more important than marginal improvements in placement quality.",
        "This conceptual framework therefore bridges the literature review and the methodology by providing the logic for the empirical comparison. It clarifies what will be compared, why the comparison matters, and how the results will be interpreted in terms of practical method selection for multi-site SDN controller placement.",
    ]),
    ("8. Research Methodology", [
        "This study adopts a methodology designed to produce objective, reproducible, and decision-relevant evidence on the comparative performance of AI-driven and heuristic controller placement methods in multi-site SDN environments. Each methodological choice is selected to support direct comparison across latency, reliability, and computational cost, while preserving experimental control and traceability.",
    ]),
    ("8.1 Research Philosophy", [
        "This study is grounded in a positivist research philosophy. Positivism is appropriate because the phenomenon under investigation is measurable and testable: controller placement methods produce observable outcomes such as latency, reliability, runtime, and resource cost. The aim of the study is not to interpret subjective experiences, but to evaluate whether one algorithmic class performs better than another under controlled conditions. A positivist stance therefore fits the research problem because it treats performance as an objective reality that can be measured, compared, and replicated.",
        "The choice of positivism is also justified by the need for evidence that can support defensible method selection in practical SDN environments. Since the central question is whether AI-based complexity is justified relative to heuristic simplicity, the study requires quantifiable results rather than interpretive narratives. Positivism enables the use of controlled simulation, repeated trials, and statistical comparison, all of which strengthen reliability and reduce the influence of researcher bias.",
    ]),
    ("8.2 Research Approach", [
        "The study follows a deductive research approach. Deduction is suitable because the investigation begins with existing ideas from optimisation theory and controller placement research, then tests them empirically through experiment. In particular, the study tests the proposition that AI-based methods may outperform heuristic methods in some scenarios, but not necessarily in all, once computational cost and reliability are included alongside latency.",
        "A deductive approach is justified because the research is not trying to generate theory from open-ended field observation. Instead, it is evaluating predefined methods against explicit performance criteria. This aligns with the golden thread of the proposal: the problem statement identifies a decision gap, the research questions ask whether AI methods offer meaningful advantages, and the methodology tests that claim through controlled experimentation. Deduction is therefore the correct approach because it moves logically from theory and prior claims to empirical verification.",
    ]),
    ("8.3 Research Design", [
        "The study uses an experimental research design. Experimental design is appropriate because the core purpose is to compare method performance under controlled and repeatable conditions. By manipulating the controller placement method and systematically varying topology and scale, the study can isolate the effect of each factor on the performance metrics of interest.",
        "A factorial design is used within the experimental framework because the study must examine more than one independent variable at the same time. In this case, the key variables include method type, topology family, and network scale. A factorial design is justified because it enables the researcher to observe not only the main effect of each variable, but also interaction effects. For example, it is important to determine whether an AI method performs better only on certain topology types, or whether its advantage changes as network size increases. These interaction effects are central to the study’s objectives and cannot be captured adequately by a single-factor design.",
    ]),
    ("8.4 Data Collection", [
        "Data collection in this study is conducted in a technical simulation environment rather than through human participants. This is appropriate because the research problem concerns algorithmic performance, not perception or behaviour. The data collection process is therefore designed to preserve experimental control, ensure repeatability, and minimise the risk of contamination from uncontrolled external factors.",
    ]),
    ("8.4.1 Design of Instruments", [
        "In this study, the instruments are the technical components used to generate, capture, and process experimental evidence. These include the simulation environment, controller placement implementations, custom Python scripts for running experiments, logging mechanisms, and metric extraction utilities. The simulation environment functions as the primary instrument because it allows the study to model multi-site SDN topologies under consistent conditions. The algorithm implementations serve as operational instruments for executing the comparative methods. The scripts and monitoring probes serve as measurement instruments for capturing latency, reliability outcomes, and computational cost.",
        "This design is justified because simulation-based instrumentation offers repeatability, parameter control, and safety. It also avoids negative operational impact on live network systems, which is important for ethical and methodological reasons. Since the study does not rely on production infrastructure, there is no risk of disrupting real network performance or compromising external systems. In addition, the measurement pipeline is designed to preserve data integrity by using structured logging, consistent metric definitions, and repeated trial execution under identical conditions.",
    ]),
    ("8.4.2 Data Collection Procedures", [
        "Data collection will proceed in a controlled sequence. First, the experimental topologies will be generated or selected according to the topology families and scale levels defined in the design. Second, each controller placement method will be executed under the same topology and scale conditions so that comparisons are fair and directly attributable to the method under test. Third, each run will capture the relevant output files, including simulation logs and packet-level traces where necessary, to support metric extraction.",
        "Latency will be derived from the time taken for control-related communication or decision propagation within the simulated environment. Reliability will be measured using outcomes such as controller coverage, failure tolerance, or service continuity under predefined failure conditions. Computational cost will be extracted from runtime, convergence duration, and related execution overhead. These metrics are selected because they jointly represent the performance trade-off that the research questions are designed to evaluate.",
        "The collection procedure is justified by the need for consistency and comparability. Running all methods under the same controlled conditions ensures that performance differences are not caused by differences in topology generation, simulation timing, or measurement procedure. Capturing both logs and packet-level data also increases traceability, allowing the experiment to be audited and repeated if needed. This is consistent with the study’s emphasis on reproducibility and transparent evidence.",
    ]),
    ("8.4.3 Piloting of Instruments", [
        "Before the full experiment begins, a pilot run or smoke test will be conducted. The purpose of the pilot is to verify that the simulation environment is stable, that the scripts execute correctly, that the logging pipeline captures the intended outputs, and that the chosen metrics can be extracted without data loss or corruption. A pilot is especially important in a computational study because small configuration errors can produce misleading results across a full factorial design.",
        "The pilot is justified on methodological grounds because it reduces the risk of invalid or incomplete data in the main experiment. It also allows any issues with runtime, file generation, or metric parsing to be corrected early, before the full experimental matrix is executed. In addition, the pilot helps confirm that the chosen network sizes and topology settings are computationally feasible within the available resources.",
    ]),
    ("8.5 Data Analysis", [
        "The data analysis strategy combines descriptive, inferential, and multi-objective techniques. Descriptive statistics will first be used to summarise the results for each method and scenario, including means, dispersion, and performance ranges. This provides an initial picture of how the methods behave across conditions and supports transparency in reporting.",
        "Inferential statistics will then be used to compare methods more formally. Where the experimental structure involves two methods under comparable conditions, t-tests may be used to assess whether observed differences are statistically meaningful. Where multiple factors are involved, analysis of variance is more appropriate because it can test the effect of method type, topology, scale, and their interactions in one model. This is justified because the research design is factorial and the study aims to understand not only whether methods differ, but also under what conditions those differences emerge.",
        "Pareto front analysis will be used to evaluate multi-objective trade-offs across latency, reliability, and computational cost. This method is appropriate because the study does not define success using a single metric. Instead, it seeks to identify whether a method dominates others across the combined objective space, or whether each method occupies a different trade-off region. Pareto analysis therefore aligns directly with the research aim and provides a rigorous basis for judging whether AI complexity is justified relative to heuristic simplicity.",
        "This combined analysis strategy is justified because it matches the nature of the research problem. Descriptive statistics provide clarity, inferential tests provide statistical support, and Pareto analysis provides decision insight. Together, they allow the study to move from raw performance data to a defensible comparison and then to practical guidance for controller placement selection.",
    ]),
    ("9. Research Quality: Ethical and Safety Issues", []),
    ("9.1 Research Quality", []),
    ("9.1.1 Reliability and Validity", [
        "This study will ensure internal validity through strict control of simulation variables so that observed differences in performance can be attributed to the controller placement method rather than to uncontrolled experimental noise. The same topology families, scale levels, traffic assumptions, and evaluation procedures will be applied consistently across all experimental runs. Method comparisons will therefore be made under equivalent conditions, allowing a credible assessment of whether AI-driven methods genuinely outperform heuristic baselines.",
        "Reliability will be strengthened through reproducibility controls. These include the use of fixed random seeds, standardised topology datasets such as Internet2- or GÉANT-type network structures, and repeatable simulation scripts. If the experiment is repeated under the same configuration, the results should be broadly consistent, subject to normal stochastic variation. This approach supports both repeatability and transparency, which are essential for a computational study of algorithmic performance.",
    ]),
    ("9.2 Ethical and Safety Considerations", []),
    ("9.2.1 Informed Consent", [
        "This study does not involve human participants, survey respondents, or interviewees. It is an algorithmic evaluation conducted entirely within simulated network environments. As a result, traditional informed consent procedures do not apply.",
    ]),
    ("9.2.2 Debriefing", [
        "Debriefing is not applicable because no human participants are exposed to the study procedures. The outputs of the study are simulation results, performance metrics, and comparative analyses of algorithms.",
    ]),
    ("9.2.3 Withdrawal from Investigation", [
        "Participant withdrawal does not apply because no participants are recruited. However, experimental runs may be stopped, repeated, or discarded if they are found to be corrupted, unstable, or methodologically invalid.",
    ]),
    ("9.2.4 Protection of Data and Systems", [
        "The study is ethical because it is conducted in a sandboxed simulation environment, such as Mininet and Python-based experimental scripts, and therefore poses no risk to live university network infrastructure, operational systems, or human well-being. The work is limited to controlled computational experiments and does not interfere with production services or expose real users to harm.",
    ]),
    ("9.2.5 Confidentiality and Anonymity", [
        "Confidentiality concerns are minimal because the study will use only public or synthetic topology data and will not require access to sensitive institutional network configurations. No private network identifiers, user data, or operational logs from live systems will be collected or disclosed. This ensures that the research remains ethically appropriate and does not compromise institutional confidentiality.",
    ]),
    ("9.2.6 Data Storage", [
        "Research data, experimental logs, scripts, and generated outputs will be stored securely in password-protected local environments during active analysis. Backup copies will be maintained in encrypted cloud storage approved by the University, such as OneDrive or Google Drive where available. Version control will be managed through a private GitHub or Bitbucket repository to preserve a clear audit trail of code and analysis changes. Access to all files will be restricted to the researcher and, where required, the supervisor. This storage plan supports data integrity, traceability, and responsible research practice.",
        "Overall, this study is ethical because it adheres to the principles of research integrity, transparency, and computational safety. It involves no human participants, does not expose sensitive systems, and is designed to generate reproducible evidence within a controlled and low-risk technical environment.",
    ]),
    ("10. Intellectual Property", [
        "This research is conducted under the auspices of the University of Zululand, and all intellectual property generated during the study shall belong to the University in accordance with institutional policy. This includes the research design, analytical framework, experimental artefacts, scripts, and any original outputs produced for the project.",
        "The source code developed for the ai-sdn-controller-placement project will be documented and managed in line with the University’s research integrity guidelines. All code contributions, simulation assets, and analytical outputs will be version-controlled, appropriately attributed, and retained in a manner that supports transparency, reproducibility, and academic accountability. Where external datasets, libraries, or tools are used, they will be acknowledged in accordance with applicable licensing and citation requirements.",
    ]),
    ("11. Resources Required and Project Plan", []),
    ("11.1 Resources Required", [
        "This study requires a modest set of technical resources because it is a software-based and simulation-driven investigation rather than a field study.",
        "The primary resource is a reliable development machine, preferably a laptop or desktop with at least a quad-core processor, 16 GB RAM, and sufficient storage space for simulation logs, datasets, and version-controlled project files. A larger memory capacity is desirable because repeated simulation runs and analysis scripts may produce substantial log output. A stable internet connection is also required for literature access, package installation, and repository synchronisation.",
        "The software environment will consist of Python 3.x together with the libraries required to implement, run, and analyse the experiments. These include standard scientific and data-processing libraries such as NumPy, Pandas, SciPy, and Matplotlib, as well as simulation and network-modelling tools such as Mininet where applicable. If the experimental workflow includes custom orchestration or metric extraction, additional Python utilities may be used to support automation, parsing, and visualisation. For project management and reproducibility, the source code will be maintained in the ai-sdn-controller-placement repository using version control.",
        "Cloud storage will also be used for secure backup and synchronisation of scripts, results, and analysis outputs. This ensures continuity of work, prevents accidental data loss, and supports traceability across the project lifecycle.",
    ]),
    ("11.2 Budget", []),
    ("11.3 Project Plan", []),
    ("12. Feasibility of the Study", [
        "The study is feasible because it is technically manageable, methodologically contained, and aligned with the researcher’s current capabilities and resources. The researcher has a background in Computer Science and Mathematics, which provides the necessary foundation for understanding algorithmic design, simulation logic, performance evaluation, and statistical comparison. This background is appropriate for a study that combines network optimisation, experimental analysis, and quantitative reasoning.",
        "The study is also feasible because the required resources are readily available. The technical environment is based on a personal development machine, Python, open-source libraries, and simulation tools such as Mininet, all of which can be used without specialised infrastructure. The project does not depend on human participants, physical network deployment, or costly proprietary software, which reduces both complexity and risk.",
        "The study is further feasible because the timeline is realistic for an Honours submission. The work is divided into clearly defined stages covering review, setup, experimentation, analysis, and writing. This sequencing allows progress to be measured and adjusted if needed, while still keeping the project within a single academic cycle.",
        "The study is feasible because the scope is controlled, the methods are reproducible, the computational requirements are moderate, and the expected outputs are achievable within the available time frame. The project therefore satisfies the practical and academic conditions necessary for successful completion at Honours level.",
    ]),
    ("13. Dissemination Plan", [
        "The findings of this study will be disseminated in two main ways. First, the finalised source code, simulation scripts, and supporting artefacts will be shared with the wider open-source community through GitHub as a form of research reciprocity. This will support transparency, reproducibility, and reuse by other researchers working on SDN optimisation and related network intelligence problems.",
        "Second, the results will be prepared for presentation at an appropriate computer science forum, such as the South African Institute of Computer Scientists and Information Technologists (SAICSIT) conference, or at an institutional research day at the University of Zululand. This will provide an opportunity to share the study’s methods, results, and practical implications with academic peers and supervisors. Dissemination in these settings also strengthens the visibility of the research and supports broader scholarly engagement.",
    ]),
    ("14. Proposed Chapter Division", [
        "Chapter 1: Introduction and Background. This chapter introduces the study, provides the research background, defines the problem statement, and presents the research questions, aim, objectives, and significance.",
        "Chapter 2: Literature Review and Conceptual Framework. This chapter reviews the empirical and theoretical literature on SDN controller placement and presents the conceptual framework that guides the study.",
        "Chapter 3: Research Methodology. This chapter explains the research philosophy, approach, design, instruments, data collection procedures, and analysis methods.",
        "Chapter 4: Presentation of Results and Data Analysis. This chapter presents the experimental results, statistical analysis, and trade-off evaluation of the controller placement methods.",
        "Chapter 5: Summary, Conclusion, and Recommendations. This chapter summarises the findings, states the conclusions, and offers recommendations for future research and practical implementation.",
    ]),
    ("15. References", []),
]

budget_rows = [
    ["Item", "Description", "Estimated Cost"],
    ["Development machine", "Personal laptop/desktop", "R0.00"],
    ["Python and libraries", "Open-source software stack", "R0.00"],
    ["Mininet and related tools", "Open-source simulation tools", "R0.00"],
    ["Cloud storage", "University-provided or free-tier backup space", "R0.00"],
    ["Internet access", "Existing personal or institutional access", "R0.00"],
    ["Printing and binding", "Final submission materials, if required", "To be determined"],
    ["Total", "", "R0.00 excluding optional printing"],
]

project_rows = [
    ["Phase", "Activities", "Indicative Timing"],
    ["Literature Review and Proposal", "Refine topic, review empirical literature, finalise proposal sections", "Current phase"],
    ["Environment Setup and Scripting", "Configure Python environment, build simulation scripts, define metrics and logging pipeline", "Next phase"],
    ["Data Collection", "Run simulations, execute repeated trials, capture logs and outputs", "Following setup"],
    ["Data Analysis and Results", "Process data, compute statistics, compare methods, generate figures and tables", "After data collection"],
    ["Final Dissertation Writing and Submission", "Integrate findings, refine discussion, complete final formatting and submission", "Final phase"],
]


def write_docx():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)

    for heading, paras in sections:
        p = doc.add_paragraph()
        r = p.add_run(heading)
        r.bold = True
        r.font.size = Pt(12)
        if paras:
            for text in paras:
                doc.add_paragraph(text)
        if heading == "11. Resources Required and Project Plan":
            doc.add_paragraph("11.2 Budget", style=None).runs[0].bold = True
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, cell in enumerate(hdr):
                cell.text = budget_rows[0][i]
            for row in budget_rows[1:]:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val
            doc.add_paragraph("11.3 Project Plan", style=None).runs[0].bold = True
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, cell in enumerate(hdr):
                cell.text = project_rows[0][i]
            for row in project_rows[1:]:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val
        if heading == "15. References":
            break

    doc.save(DOCX_PATH)


def citation_link_markup(text):
    def _replace(match):
        ref_id = match.group(1)
        if ref_id in REFERENCE_IDS:
            return f'<a href="#ref{ref_id}"><u>[{ref_id}]</u></a>'
        return match.group(0)

    return CITATION_PATTERN.sub(_replace, text)


def pdf_paragraph(text, style, with_links=False):
    value = citation_link_markup(text) if with_links else text
    return Paragraph(value.replace("&", "&amp;"), style)


def write_pdf():
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        alignment=1,
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        spaceBefore=8,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        spaceAfter=6,
    )
    small_style = ParagraphStyle(
        "SmallStyle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        spaceAfter=4,
    )

    story = [pdf_paragraph(TITLE, title_style), Spacer(1, 4)]
    for heading, paras in sections:
        story.append(pdf_paragraph(heading, heading_style))
        if heading == "11. Resources Required and Project Plan":
            for text in [
                "11.1 Resources Required",
                *sections[[h for h, _ in sections].index(heading)][1],
                "11.2 Budget",
            ]:
                pass
        for text in paras:
            story.append(pdf_paragraph(text, body_style, with_links=True))
        if heading == "11. Resources Required and Project Plan":
            story.append(pdf_paragraph("11.2 Budget", heading_style))
            budget_table = Table(budget_rows, colWidths=[42*mm, 98*mm, 35*mm])
            budget_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(budget_table)
            story.append(Spacer(1, 6))
            story.append(pdf_paragraph("11.3 Project Plan", heading_style))
            project_table = Table(project_rows, colWidths=[42*mm, 98*mm, 35*mm])
            project_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(project_table)
            story.append(Spacer(1, 8))
        if heading == "15. References":
            for ref_id, ref_text, ref_url in REFERENCE_ITEMS:
                ref_line = f'<a name="ref{ref_id}"/>[{ref_id}] {ref_text} <a href="{ref_url}">{ref_url}</a>'
                story.append(pdf_paragraph(ref_line, small_style))
            break

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
    )
    doc.build(story)


if __name__ == "__main__":
    write_docx()
    write_pdf()
    shutil.copy2(PDF_PATH, TEMPLATE_PDF_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {TEMPLATE_PDF_PATH}")
